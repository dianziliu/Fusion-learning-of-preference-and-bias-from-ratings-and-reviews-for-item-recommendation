import json
import os
import re
import sys
import time

import numpy as np
import tensorflow as tf
from docx import Document
from docx.shared import Inches, Pt

from .data_loader import Data_Loader
from .utils import clean_str


def readfile(filename):
	f = open(filename, encoding='utf-8')
	data = []
	for line in f.readlines():
		line = json.loads(line)
		data.append(line)
	f.close()
	return data

def visual(data, uit, data_loader, utexts, itexts, u_texts, i_texts, filename):
    #idx: index of the sampled data point in the data_loader.train_uit
    word2idx = data_loader.word2idx
    idx2word = {v[1]: v[0] for v in word2idx.items()}

    raw_data = readfile(filename)

    document = Document()

    for idx, (doc_idx, doc) in enumerate(zip(utexts, u_texts)):
        attention_score = data[2][-1][idx]
        document.add_paragraph('attention score: '+str(attention_score))
        word_atts = data[0][idx]
        raw_doc = raw_data[doc_idx-1]

        visual_single_doc(word_atts, doc, doc_idx, idx2word,
                          word2idx, raw_doc, document)

    document.add_paragraph(
        '-----------------------------------------------------')
    for idx, (doc_idx, doc) in enumerate(zip(itexts, i_texts)):
        attention_score = data[3][-1][idx]
        document.add_paragraph('attention score: '+str(attention_score))
        word_atts = data[1][idx]
        raw_doc = raw_data[doc_idx-1]

        visual_single_doc(word_atts, doc, doc_idx, idx2word,
                          word2idx, raw_doc, document)

    document.save('_'.join(map(str, uit))+'_atts.docx')
    doc_atts = []
    for user_layer in data[2]:
        doc_att = np.squeeze(user_layer)
        doc_atts.append(doc_att)
    for item_layer in data[3]:
        doc_att = np.squeeze(item_layer)
        doc_atts.append(doc_att)
    np.savetxt('_'.join(map(str, uit))+'_doc_atts.txt', doc_atts, fmt='%.5f')
    # doc_atts = [round(value,5) for value in doc_atts]
    # print(' '.join(map(str,doc_atts)))

    #data[0]: user word-level attention [1,6,60]
    #data[1]: item word-level attention
    #data[3]: user document-level attention list of [1,6,1]
    #data[4]: item document-level attention
    pass


def visual_single_doc(attentions, word_vec, doc_idx, idx2word, word2idx, raw_doc, document):
    print('text id: ', doc_idx)
    import nltk
    raw_tokens = nltk.word_tokenize(
        clean_str((raw_doc['reviewText']+' '+raw_doc['summary']).lower()))
    idx = len(raw_tokens)-1
    counter = 0

    while counter < 80 and idx > 0:
        if raw_tokens[idx] in word2idx:
            counter += 1
        idx = idx - 1
    raw_tokens = raw_tokens[idx:]

    word2att = {}
    res = []
    for att, word_idx in zip(attentions, word_vec):
        if word_idx != 0:
            word = idx2word[word_idx]
            att = round(att, 5)
            word2att[word] = max(word2att.get(word, 0), att)

    atts = []
    for token in raw_tokens:
        att = word2att.get(token, 0)
        # res.append(token+':'+str(att))
        atts.append(att)
    single_document(raw_tokens, atts, document)
    # print(' '.join(res))


def single_document(tokens, atts, document):
    atts = np.array(atts)

    nz_atts = atts[atts != 0]
    min_size = 12
    max_size = 20
    if len(nz_atts) < 2:
        return

    min_att = min(nz_atts)
    max_att = max(nz_atts)

    if min_att == max_att:
        new_atts = np.array([15]*len(nz_atts))
    else:
        new_atts = 12+(max_size - min_size)/(max_att - min_att)*(nz_atts - min_att)

    atts[atts != 0] = new_atts
    atts = atts.astype(np.int32)

    p = document.add_paragraph('')

    for token, att in zip(tokens, atts):
        run = p.add_run(token+' ')
        if att > 0:
            run.bold = True
            run.font.size = Pt(att)


def visualization(flags, model, data_loader, filename):

    # 1. 加载模型


    # 2. 准备数据
    # checkpoint_path=os.path.join(model.ckpt_dir,"{}.h5".format(model.name))
    # Big_model=model
    # model=model.model
    # model.load_weights(checkpoint_path)
    model=model.model
    # user,item,label, utexts, itexts, text= data_loader.sample_point()
    user, item, label, utexts, itexts, text = data_loader.find_a_user()

    utexts = utexts.astype(int)
    itexts = itexts.astype(int)

    feed_dict={"u_input": user,
                      "i_input": item,
                      "text": text,
                      "utext": utexts,
                      "itext": itexts,
                      "label":label
    }
    # 需要用到以下4个值，因此在可视化的时候需要对模型的输出进行一定的修改。
    # model.word_user_alpha, model.word_item_alpha, model.doc_user_alpha, model.doc_item_alpha
    res = model.predict(feed_dict)
    res=res[1:]
    print(utexts.dtype)
    u_texts = data_loader.vec_texts[utexts]
    i_texts = data_loader.vec_texts[itexts]

    res[2] = np.array(res[2]).transpose(1, 0, 2)
    res[3] = np.array(res[3]).transpose(1, 0, 2)

    for i in range(len(user)):
        uit = [user[i], item[i], label[i]]
        print(uit)
        res_trans = []
        for r in res:
            res_trans.append(r[i])
        visual(res_trans, uit, data_loader,
               utexts[i], itexts[i], u_texts[i], i_texts[i], filename)
